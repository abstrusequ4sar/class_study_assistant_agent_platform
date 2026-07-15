"""将逐页 Markdown 演讲稿转换为便于打印和排练的 Word 文档。"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "speech.md"
OUTPUT = ROOT / "课程学习助手Agent平台新版汇报演讲稿.docx"
FONT = "Microsoft YaHei"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run._r.append(field)


def build():
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    add_page_number(section.footer.paragraphs[0])

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.35
    normal.paragraph_format.space_after = Pt(7)

    title_style = document.styles["Title"]
    title_style.font.name = FONT
    title_style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    title_style.font.size = Pt(22)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(31, 78, 121)

    heading = document.styles["Heading 1"]
    heading.font.name = FONT
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    heading.font.size = Pt(16)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor(31, 78, 121)
    heading.paragraph_format.space_after = Pt(10)

    slide_started = False
    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(line[2:])
            continue
        if line.startswith("## "):
            if slide_started:
                document.add_page_break()
            slide_started = True
            document.add_paragraph(line[3:], style="Heading 1")
            continue
        if line.endswith("  "):
            line = line[:-2]
        if line.startswith("建议用时："):
            paragraph = document.add_paragraph()
            run = paragraph.add_run(line)
            set_run_font(run, 10, True, (112, 173, 71))
            paragraph.paragraph_format.space_after = Pt(10)
            continue
        if len(line) > 2 and line[0].isdigit() and ". " in line[:4]:
            paragraph = document.add_paragraph(style="List Number")
            paragraph.add_run(line.split(". ", 1)[1])
            continue
        paragraph = document.add_paragraph()
        paragraph.add_run(line.replace("`", ""))

    document.core_properties.title = "课程学习助手 Agent 平台逐页演讲稿"
    document.core_properties.subject = "大型程序设计实践汇报"
    document.core_properties.author = "课程学习助手 Agent 平台项目组"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
